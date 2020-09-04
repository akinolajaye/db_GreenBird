import tkinter as tk #imports the tkinter module which allow gui programming
from tkinter import ttk
from editdatbwin import EditDatabase
from tkinter import scrolledtext as st
import bookingsdatabasesql as bdbsql
import propertydatabasesql as pdbsql
import customersql as cdbsql    
from tkinter import messagebox as mbx
import customerwin as cr
import datetime
from docx import *
from docx.enum.text import WD_ALIGN_PARAGRAPH

class BookCustomers(EditDatabase): #inherits from editdatabase
    def __init__(self,win): 
        self.win=win
        self.bookings_window=tk.Toplevel(win)#creates a instance of the tkinter module in the form of a new window
        self.bookings_window.geometry('1090x440')#sets a default size for the window when opened
        self.bookings_window.resizable(0,0)
        self.createStyle()
        self.createWidgets()
        self.array=[self.name_entry,self.duration_entry,self.start_date_entry,self.end_date_entry,self.num_of_people_entry,\
            self.flat_name_entry,self.rent_due_date_entry]#creates a list that holds all entrys


    def createFrame(self):

        main_frame =ttk.LabelFrame(self.bookings_window,text ='GreenBird Properties Bookings')#creaetes a main frame that holds all frames
        main_frame.pack(anchor ='w',fill='both')

       
        self.data_frame=ttk.Frame(main_frame,style ='DF.TLabelframe')#creates a frame that holdd the data widgets and frames
        self.data_frame.pack(side=tk.TOP,fill='x',pady=6)
        
        self.data_entry_frame=ttk.LabelFrame(self.data_frame,text ='Bookings:',style ='DEF.TLabelframe')#creates a labeled frame that will hold all data entries
        self.data_entry_frame.pack(side=tk.LEFT,anchor ='nw')
        

        self.data_display_frame=ttk.LabelFrame(self.data_frame,text='Available Bookings:',style ='DEF.TLabelframe')#creates a labeled frame that will hold all displayed data
        self.data_display_frame.pack(side=tk.RIGHT,anchor ='w')
        
        self.button_frame=ttk.Frame(main_frame)#creates a button frame that will hold all buttons in the mainframe
        self.button_frame.pack(side=tk.BOTTOM,fill='x')   
        
    def createLabels(self):

        #creates a label
        name_lbl=ttk.Label(self.data_entry_frame,text='Customer Name: ',style ='DEF.TLabel')
        name_lbl.grid(row=1,column=0)

        duration_lbl=ttk.Label(self.data_entry_frame,text='Duration (Months):',style ='DEF.TLabel')
        duration_lbl.grid(row=2,column=0)

        start_date_lbl=ttk.Label(self.data_entry_frame,text='Start Date (dd/mm/yyyy):',style ='DEF.TLabel')
        start_date_lbl.grid(row=3,column=0)

        end_date_lbl=ttk.Label(self.data_entry_frame,text='End Date:',style ='DEF.TLabel')
        end_date_lbl.grid(row=4,column=0)
        

        num_of_people_lbl=ttk.Label(self.data_entry_frame,text='Number of People: ',style ='DEF.TLabel')
        num_of_people_lbl.grid(row=5,column=0)

        flat_name_lbl=ttk.Label(self.data_entry_frame,text='Flat Name: ',style ='DEF.TLabel') #creates a label
        flat_name_lbl.grid(row=6,column=0)

        rent_due_date_lbl=ttk.Label(self.data_entry_frame,text='Rent Due Date: ',style ='DEF.TLabel')
        rent_due_date_lbl.grid(row=7,column=0)

        tday=datetime.datetime.now() #uses the date time module to get current date
        date=self.getDateString(tday)#converts the date into a string

        date_lbl=ttk.Label(self.data_entry_frame,text = f'{date}')
        date_lbl.grid(row=8,column=0,sticky='w')

    def createEntrys(self):

        name=tk.StringVar()# declares variables types
        flat_name=tk.StringVar()
        duration=tk.IntVar()
        start_date=tk.StringVar()
        end_date=tk.StringVar()
        num_of_people=tk.IntVar() 
        
        rent_due_date=tk.StringVar() 
        phone_number=tk.StringVar()
        email=tk.StringVar()


        #creates entries
        valid =self.data_entry_frame.register(self.intEntryValidation) #same as the edit database class

        self.name_entry=ttk.Combobox(self.data_entry_frame,textvariable =name,width=40,font =('calibri',15),postcommand=self.autoFillNames)
        self.name_entry.grid(row=1,column=1) 


        self.duration_entry=ttk.Entry(self.data_entry_frame,textvariable =duration,width=40,font =('calibri',15),\
            validate ='key',validatecommand=(valid,'%P'))
        self.duration_entry.grid(row=2,column=1)

        self.start_date_entry=ttk.Entry(self.data_entry_frame,textvariable =start_date,width=40,font =('calibri',15))
        self.start_date_entry.grid(row=3,column=1) 

        self.end_date_entry=ttk.Combobox(self.data_entry_frame,textvariable =end_date,width=40,font =('calibri',15),\
            postcommand= lambda: self.autoFillEndDate(self.start_date_entry.get()))
        self.end_date_entry.grid(row=4,column=1) 


        self.num_of_people_entry=ttk.Entry(self.data_entry_frame,textvariable =num_of_people,width=40,font =('calibri',15),\
            validate ='key',validatecommand=(valid,'%P'))
        self.num_of_people_entry.grid(row=5,column=1)

        self.flat_name_entry=ttk.Combobox(self.data_entry_frame,textvariable =flat_name,width=40,font =('calibri',15)\
            ,postcommand= self.autoFillFlatName) 
        self.flat_name_entry.grid(row=6,column=1)

        self.rent_due_date_entry=ttk.Combobox(self.data_entry_frame,textvariable =rent_due_date,width=40,font =('calibri',15),\
            postcommand=self.autoFillPayDate)
        self.rent_due_date_entry.grid(row=7,column=1)  
       
    def createWidgets(self):
      
        self.createFrame()
        self.createLabels()
        self.createEntrys()
        self.createListBox(self.data_display_frame,100,12)
        self.listboxCommand()
        self.createButtons()

        for i in self.data_entry_frame.winfo_children():
            i.grid_configure(pady=5,sticky='w') 
        
    def listboxCommand(self):

        self.display_box.bind('<<ListboxSelect>>',lambda event :self.recordDisplayIndex(event,self.array))

    def createButtons(self):

        #creates click buttons
        add_customerb=ttk.Button(self.data_entry_frame,text='Add New Customer',command =self.openCustomerWindow)
        add_customerb.grid(row=0,column=0)

        add_datab=ttk.Button(self.button_frame,text='Book Customer',command =self.addBookingsData)
        add_datab.grid(row=0,column=0)

        search_datab=ttk.Button(self.button_frame,text='Search Customer',command=self.searchBookingsDatabase)
        search_datab.grid(row=0,column=1)

        display_datab=ttk.Button(self.button_frame,text='Display Customers',command=lambda:self.displayTable('Bookings'))
        display_datab.grid(row=0,column=2)

        clear_datab=ttk.Button(self.button_frame,text='Clear',command=lambda:self.clearData(self.array))
        clear_datab.grid(row=0,column=3)

        delete_datab=ttk.Button(self.button_frame,text='Delete',command =self.delete)
        delete_datab.grid(row=0,column=4)

        self.update_datab=ttk.Button(self.button_frame,text='Update',command =self.updateBookingDatabase)
        self.update_datab.grid(row=0,column=5)

        exitb=ttk.Button(self.button_frame,text='Exit',command=lambda:self._exit(self.bookings_window))
        exitb.grid(row=0,column=6)

        for i in self.button_frame.winfo_children():
            i.configure(style= 'B.TButton')

    def bookingValid(self,update):

        valid=False
        empty=True

        empty=self.emptyEntryCheck(8)

        if empty:
            mbx.showerror('Error','Fill out all remaining fields!')
            return valid,empty

        if self.BookingAvailabiltyValidate(update):
            valid=True
        else:
            valid=False
            mbx.showerror('Error',f'{self.flat_name_entry.get()} has been leased already!')
            return valid,empty

        if self.validateDate(self.start_date_entry.get()):
            valid=True
        else:
            valid=False
            mbx.showerror('Error','Incorrect format for date should be DD/MM/YYYY!')
            return valid,empty

        if self.validateDate(self.end_date_entry.get()):
            valid=True
        else:
            valid=False
            mbx.showerror('Error','Incorrect format for date should be DD/MM/YYYY!')
            return valid,empty

        return valid,empty

    def BookingAvailabiltyValidate(self,update):

        flats=bdbsql.getAllFlats()#gets all flats 
        if self.flat_name_entry.get() in flats and not update: #checks if the flat has been booked

            return False
        else: 

            return True

    def validateDate(self,date_text):
        try:
            date_text=date_text.split('/')
            date=date_text[0]+'-'+date_text[1]+'-'+date_text[2]
            datetime.datetime.strptime(date, '%d-%m-%Y')
            return True
        except:

            return False 

    def emptyEntryCheck(self,entry_start_index):
        entrys=self.data_entry_frame.winfo_children()
        
        for i in range(entry_start_index,len(entrys)-1): #starts at 5 because widgets before that are labels and we only want entry widgets
            if entrys[i].get() =='': #checks if an entry is empty
                empty=True
                return empty
            else:
                empty=False

        return empty

    def addBookingsData(self):

        valid,empty=self.bookingValid(False)


        if valid and not empty:
            bdbsql.addBookingsData(self.duration_entry.get(),self.start_date_entry.get(),\
                self.end_date_entry.get(),self.num_of_people_entry.get(),\
                self.flat_name_entry.get(),self.rent_due_date_entry.get()) #adds data from entrys and puts in sql database
            self.addCustBooking()
            self.displayTable('Bookings')
            self.createDocument() #creates a document based on bookings

    def getSqlCommand(self):

        sql_string=''
        sql_variables=[]

        if self.name_entry.get().isspace() or self.name_entry.get()=='':
            pass
        else:
            first_name,surname=self.splitName(self.name_entry.get())
            sql_string+='first_name=? AND surname=? '

            sql_variables.append(first_name)
            sql_variables.append(surname)


        if self.duration_entry.get().isspace() or self.duration_entry.get()=='':
            pass
        else:
            if sql_string=='':
                sql_string+='duration=? '
            else:
                sql_string+='AND duration=? '

            sql_variables.append(self.duration_entry.get())

        if self.start_date_entry.get().isspace() or self.start_date_entry.get()=='':
            pass
        else:
            if sql_string=='':
                sql_string+='start_date=? '
            else:
                sql_string+='AND start_date=? '

            sql_variables.append(self.start_date_entry.get())

        if self.end_date_entry.get().isspace() or self.end_date_entry.get()=='':
            pass
        else:
            if sql_string=='':
                sql_string+='end_date=? '
            else:
                sql_string+='AND end_date=? '
            sql_variables.append(self.end_date_entry.get())

        if self.num_of_people_entry.get().isspace() or self.num_of_people_entry.get()=='':
            pass
        else:
            if sql_string=='':
                sql_string+='num_of_people=? '
            else:
                sql_string+='AND num_of_people=? '
            sql_variables.append(self.num_of_people_entry.get())

        if self.flat_name_entry.get().isspace() or self.flat_name_entry.get()=='':
            pass
        else:
            if sql_string=='':
                sql_string+='flat_name=? '
            else:
                sql_string+='AND flat_name=? '
            sql_variables.append(self.flat_name_entry.get())


        if self.rent_due_date_entry.get().isspace() or self.rent_due_date_entry.get()=='':
            pass
        else:
            if sql_string=='':
                sql_string+='rent_due_date=? '
            else:
                sql_string+='AND rent_due_date=? '

            sql_variables.append(self.rent_due_date_entry.get())


        return sql_string,sql_variables

    def searchBookingsDatabase(self):
        self.display_box.delete(0,tk.END)
        #uses entry values as search criteria

        sql_string,sql_variables=self.getSqlCommand()

        for i in bdbsql.searchBookingsData(sql_string,sql_variables):
            self.display_box.insert(tk.END,i) #inserts different data into display box based on index
     
    def createDocument(self):
        
        location_list=pdbsql.getLocation(self.flat_name_entry.get()) #gets a list that has location data of an apartment from sql database
        occpts=bdbsql.getNumOccpts(self.flat_name_entry.get()) #gets the number of occupants from sql database
        doc=Document()#creates an instance of document module using docx
        pic=doc.add_picture('logo.png',width=shared.Cm(9.04),height=shared.Cm(5.69)) #adds picture to the document
        lastpara=doc.paragraphs[-1]#puts the cursor at the last paragraph (in this case will be the pic)
        lastpara.alignment=WD_ALIGN_PARAGRAPH.CENTER #center aligns the picture

        lease_declaration_paragraph=doc.add_paragraph(f"This is an agreement to sublet real property according to the terms specified below, \
        hereinafter known as the 'Agreement'. The Sublessor, known as Adrian Akinola agrees to sublet to Subtenant, known as ")
        lease_declaration_paragraph.add_run(self.name_entry.get()[1:])
        location_para=doc.add_paragraph("The location of the premises is located at ")
        for i in location_list: #prints out the whole location on the document
            location_para.add_run(f'\n {i}')

        
        duration=doc.add_paragraph('Lease Duration: ')
        duration.add_run('\n'+str(self.duration_entry.get()))
        duration.add_run(' months')
        duration.add_run('\nStart Date: ')
        duration.add_run(self.start_date_entry.get())
        duration.add_run('\nEnd Date: ')
        duration.add_run(self.end_date_entry.get())
        
        rent=doc.add_paragraph('The Rent is £')
        rent.add_run(str(bdbsql.getRentPrice(self.flat_name_entry.get())))
        rent.add_run(' to be paid on the ')
        rent.add_run(self.rent_due_date_entry.get())
        rent.add_run('.')

        num_of_people=doc.add_paragraph('No. Of Occupants  (including Subtenant): ')
        num_of_people.add_run(str(occpts))
        
        doc.save(f'Lease Agreement for {self.name_entry.get()} at {self.flat_name_entry.get()}.docx')#saves the document

    def delete(self):


        self.deleteData('Bookings','idb')#deletes the row selected 
                
    def autoFillFlatName(self):
       
        try:
            self.flat_name_entry['values']=''
            self.flat_name_entry.set('')
            available_flats=bdbsql.getAvailableFlats()
         
            for i in available_flats:
                self.flat_name_entry['values']=(*self.flat_name_entry['values'],i)

        except:
            pass      

    def getDateString(self,date): #turns the date into a string
        
        mstr=(date.strftime('%m'))
        dstr=(date.strftime('%d'))
        ystr=(date.strftime('%Y'))
        datestr =dstr+'/'+mstr+'/'+ystr

        return datestr

    def stringDateToDatetime(self,date): #converts date string into a datetime variable to be manipulated
        datestring=date
        datestring=datestring.split('/')
        date_as_time=datetime.datetime.strptime(datestring[2]+' '+datestring[1]+' '+datestring[0],'%Y %m %d')     
        
        return date_as_time
    
    def autoFillEndDate(self,date):
        try:
            self.end_date_entry['values']='' #sets combobox to empty
            self.end_date_entry.set('') #sets the entry in combobox to empty
            months=datetime.timedelta(days=28) # creates a value that can be used to add and subtract to a date 28 days can be added to a date
            date_as_time=self.stringDateToDatetime(date) #convets date to datetime variable
            for i in range(int(self.duration_entry.get())):
                date_as_time+=months #adds 28 days depending on how many months

            date_as_time=self.getDateString(date_as_time) #converts it back to string
            self.end_date_entry['values']=date_as_time #puts it on combo box
        except:
            mbx.showerror('Error','Incorrect format for start date or duration!')
                       
    def getPayDate(self,start_date): #creates a pay date for the rent
        ten_days=datetime.timedelta(days=10)
        five_days=datetime.timedelta(days=5)
        paydate=int((start_date + ten_days).strftime('%d'))
        if paydate >28:#if the start date plus 10 days is greater than the 28th of a month
            paydate=start_date+five_days
        else:
            paydate=start_date+ten_days


        
        

        return paydate

    def autoFillPayDate(self):
        try:
            self.rent_due_date_entry['values']='' #emptys the combobox
            self.rent_due_date_entry.set('') #clears current text on combobox
            
            date=self.stringDateToDatetime(self.start_date_entry.get()) #converts date from string to datetime
            date=self.getPayDate(date) #retrieves the pay date
            date=date.strftime('%d') #gets the day from the date sting
            date=self.getEndSentence(date) #creates a sentence that is used to describe pay date (if day is 24 then will create sentence 24th of each month)
            self.rent_due_date_entry['values']=(*self.rent_due_date_entry['values'],date)#puts the whole sentence on the combobox list
            
            
        except Exception as err:
            print(err)
            
    def getEndSentence(self,day):
        import math
        day=int(day)
        ordinal = lambda n: "%d%s" % (n,"tsnrhtdd"[(math.floor(n//10)%10!=1)*(n%10<4)*n%10::4]) #gets an ordinal number for a day eg 24 returns 24th, 1 returns 1st
        day=ordinal(day)
        sentence = day +' of each month' ### add the ending month

        return sentence

    def openCustomerWindow(self):

        cr.Customer(self.win)

    def splitName(self,full_name):
        name=full_name.split(' ')
        first_name,surname=name[0],name[1]

        return first_name,surname

    def autoFillNames(self):
        try:
            self.name_entry['values']=''
            self.name_entry.set('')
            names=cdbsql.getCustomerName()

        
            for i in names:

                self.name_entry['values']=(*self.name_entry['values'],i[0])

        except Exception as err:
            print(err)
            
    def addCustBooking(self):
        cust_id=self.name_entry.get()[0]#need a way to get new id wiht same name
        cdbsql.addCustomerBooking(cust_id,self.flat_name_entry.get())#composite key


    def updateBookingDatabase(self):
        valid,empty=self.bookingValid(True)


        if valid and not empty:
            self.delete()
            self.addBookingsData()
        #updates data by inserting entry values into the the sql database
        

            
    def displayTable(self,table):######*
        self.display_box.delete(0,tk.END)#empties the list box
        for i in bdbsql.viewData():#calls sql function that returns all data values as an array
            self.display_box.insert(tk.END,i) #the value returned from the sql function is then printed on the display box
     
            









